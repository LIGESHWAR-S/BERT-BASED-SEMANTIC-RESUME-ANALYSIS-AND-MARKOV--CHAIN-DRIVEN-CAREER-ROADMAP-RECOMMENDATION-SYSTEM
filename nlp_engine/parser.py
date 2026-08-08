import re
import io
import fitz  # PyMuPDF
import docx  # python-docx

SECTION_MAPPING = {
    'summary': [r'^summary$', r'^professional summary$', r'^profile$', r'^career objective$', r'^objective$', r'^about me$', r'^executive summary$'],
    'education': [r'^education$', r'^academic background$', r'^academic qualifications$', r'^qualifications$', r'^academic history$', r'^academic profile$'],
    'experience': [r'^experience$', r'^work experience$', r'^professional experience$', r'^employment history$', r'^work history$', r'^employment$', r'^career history$'],
    'skills': [r'^skills$', r'^technical skills$', r'^core competencies$', r'^key skills$', r'^technologies$', r'^specialties$', r'^expertise$', r'^skillset$', r'^areas of expertise$'],
    'projects': [r'^projects$', r'^academic projects$', r'^personal projects$', r'^key projects$', r'^selected projects$'],
    'certifications': [r'^certifications$', r'^licenses$', r'^certificates$', r'^courses$', r'^professional certifications$'],
    'achievements': [r'^achievements$', r'^accomplishments$', r'^awards$', r'^honors$', r'^publications$'],
    'languages': [r'^languages$', r'^language proficiency$', r'^languages spoken$']
}

def clean_text(text):
    """
    Cleans raw extracted text by normalizing spaces, removing non-printable chars, etc.
    """
    if not text:
        return ""
    # Replace multiple newlines/whitespaces
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text) # Remove non-ascii control chars
    return text.strip()

def extract_text_from_pdf(file_bytes_or_path):
    """
    Extracts raw text from a PDF file using PyMuPDF (fitz).
    """
    text = ""
    if isinstance(file_bytes_or_path, bytes):
        doc = fitz.open(stream=file_bytes_or_path, filetype="pdf")
    else:
        doc = fitz.open(file_bytes_or_path)
    
    for page in doc:
        text += page.get_text() + "\n"
    
    doc.close()
    return clean_text(text)

def extract_text_from_docx(file_bytes_or_path):
    """
    Extracts raw text from a DOCX file using python-docx.
    """
    if isinstance(file_bytes_or_path, bytes):
        doc = docx.Document(io.BytesIO(file_bytes_or_path))
    else:
        doc = docx.Document(file_bytes_or_path)
    
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
                
    return clean_text("\n".join(text))

def extract_text_from_txt(file_bytes_or_path):
    """
    Extracts raw text from a TXT file.
    """
    if isinstance(file_bytes_or_path, bytes):
        try:
            return clean_text(file_bytes_or_path.decode('utf-8'))
        except UnicodeDecodeError:
            return clean_text(file_bytes_or_path.decode('latin-1'))
    else:
        with open(file_bytes_or_path, 'r', encoding='utf-8', errors='ignore') as f:
            return clean_text(f.read())

def extract_text(file_bytes, filename):
    """
    Helper to extract text based on filename extension.
    """
    ext = filename.split('.')[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_bytes)
    elif ext == 'txt':
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def identify_section_heading(line):
    """
    Checks if a line represents a known section heading.
    Returns the mapped section name or None.
    """
    cleaned = line.strip().lower()
    # Remove numbering and punctuation (e.g. "1. Education", "Work Experience:")
    cleaned = re.sub(r'^[\d\.\-\s]+', '', cleaned)
    cleaned = re.sub(r'[\:\,\.\s]+$', '', cleaned)
    cleaned = cleaned.strip()
    
    if not cleaned or len(cleaned) > 40:
        return None
        
    for section_name, patterns in SECTION_MAPPING.items():
        for pat in patterns:
            if re.match(pat, cleaned):
                return section_name
    return None

def detect_sections(text):
    """
    Splits the resume text into sections based on detected headings.
    Returns a dict with section names as keys.
    """
    lines = text.split('\n')
    sections = {
        'contact': [],
        'summary': [],
        'education': [],
        'experience': [],
        'skills': [],
        'projects': [],
        'certifications': [],
        'achievements': [],
        'languages': []
    }
    
    current_section = 'contact' # Default section at the start
    
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue
            
        detected = identify_section_heading(cleaned_line)
        if detected:
            current_section = detected
        else:
            sections[current_section].append(line)
            
    # Join lists into single strings and clean
    parsed_sections = {}
    for sec, content in sections.items():
        parsed_sections[sec] = "\n".join(content).strip()
        
    # Heuristic for Contact info: if the contact section is huge, it might mean section detection failed
    # or it is a flat resume. If so, leave contact as is, but we try to keep it reasonable.
    return parsed_sections

def extract_contact_info(text):
    """
    Extracts email, phone, and links (like github, linkedin) from text.
    """
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    url_pattern = r'(?:https?://)?(?:www\.)?(?:linkedin\.com/in/|github\.com/)[a-zA-Z0-9_-]+'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    urls = re.findall(url_pattern, text, re.IGNORECASE)
    
    return {
        'emails': list(set(emails)),
        'phones': list(set(phones)),
        'links': list(set(urls))
    }
